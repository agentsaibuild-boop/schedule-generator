"""File manager for accessing, converting, and caching local project files.

Implements Rule #0: Convert 100% of documentation BEFORE analysis.
Supports PDF (text + OCR), Excel, DOCX, CSV, JSON, TXT.
Converted files are stored in a ``converted/`` subfolder inside the
project directory together with a ``_manifest.json`` cache.
"""

from __future__ import annotations

import concurrent.futures
import csv
import io
import json
import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

_FILE_CONVERT_TIMEOUT = 120  # seconds per file

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".csv", ".json", ".txt", ".docx"}

# Minimum average characters per page to consider a PDF "text-based"
_MIN_CHARS_PER_PAGE = 50
# Под този брой знаци страницата няма годен текстов слой.
_THIN_TEXT_CHARS = 10
# Дял от площта на страницата, покрит с изображения, над който я третираме
# като сканирана — тогава решава vision, а не броят знаци (P1).
_SCANNED_IMAGE_COVERAGE = 0.5

APP_VERSION = "0.1"


class FileManager:
    """Manages access to local project files and their conversion to JSON."""

    def __init__(self, base_path: str | None = None) -> None:
        """Initialize the file manager.

        Args:
            base_path: Optional base path to the project directory.
        """
        self.base_path: Path | None = Path(base_path) if base_path else None
        self._manifest: dict[str, Any] = {}

    # ------------------------------------------------------------------
    # Project path management
    # ------------------------------------------------------------------

    def set_project_path(self, path: str) -> dict:
        """Validate and set the project directory, scanning for files.

        Args:
            path: Path to the project directory.

        Returns:
            Dict with keys: valid, files_count, converted_count, needs_conversion.
        """
        p = Path(path)
        if not (p.exists() and p.is_dir()):
            return {
                "valid": False,
                "files_count": 0,
                "converted_count": 0,
                "needs_conversion": 0,
            }

        self.base_path = p
        self._load_manifest()

        supported = self._list_supported_files()
        status = self.get_conversion_status()

        return {
            "valid": True,
            "files_count": len(supported),
            "converted_count": status["converted"],
            "needs_conversion": status["pending"] + status["changed"],
        }

    # ------------------------------------------------------------------
    # File listing helpers
    # ------------------------------------------------------------------

    @classmethod
    def looks_like_boq(cls, text: str) -> dict:
        """Дали текстът е количествена сметка — по СЪДЪРЖАНИЕ, не по име.

        Две независими улики, за да не се хване всеки документ, споменал „м3":
          1. заглавия на колони, които вървят заедно само в КСС;
          2. достатъчно на брой редове с мерна единица.

        Args:
            text: Извлеченият текст на документа.

        Returns:
            {is_boq, confidence, evidence} — `evidence` казва КОЕ го е решило,
            за да може човек да провери преценката, а не да я приема наум.
        """
        if not text or len(text) < 100:
            return {"is_boq": False, "confidence": 0.0, "evidence": []}

        lowered = text.lower()
        evidence: list[str] = []
        score = 0.0

        for markers in cls._BOQ_COLUMN_MARKERS:
            if all(marker in lowered for marker in markers):
                evidence.append("колони: " + " + ".join(sorted(markers)))
                score += 0.6
                break

        # Броим РЕДОВЕ с мерна единица, не позиции в знаци.  Групирането по
        # 200-знакови блокове недоброяваше сбити таблици — при 6 реда даваше 2.
        unit_rows = sum(
            1 for line in text.splitlines() if cls._BOQ_UNIT_RE.search(line)
        )
        if unit_rows >= cls._BOQ_MIN_UNIT_ROWS:
            evidence.append(f"{unit_rows} реда с мерна единица")
            score += 0.4
        elif unit_rows:
            evidence.append(f"само {unit_rows} реда с мерна единица")

        confidence = min(score, 1.0)
        return {
            "is_boq": confidence >= 0.6,
            "confidence": round(confidence, 2),
            "evidence": evidence,
        }

    def find_boq_by_content(self) -> dict:
        """Потърси количествена сметка сред ВЕЧЕ КОНВЕРТИРАНИТЕ документи.

        Ползва се, когато по имена не е намерен задължителен файл.  Работи
        върху `converted/`, тоест изисква конверсията да е минала.

        Returns:
            {found: [имена], details: {име: {confidence, evidence}}}
        """
        found: list[str] = []
        details: dict[str, dict] = {}

        if self.base_path is None:
            return {"found": found, "details": details}

        converted_dir = self.base_path / "converted"
        if not converted_dir.exists():
            return {"found": found, "details": details}

        for jf in sorted(converted_dir.glob("*.json")):
            if jf.name == "_manifest.json":
                continue
            try:
                data = json.loads(jf.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                continue

            source = data.get("source_file", jf.stem)
            verdict = self.looks_like_boq(data.get("full_text", ""))
            details[source] = verdict
            if verdict["is_boq"]:
                found.append(source)
                logger.info(
                    "Количествена сметка разпозната по СЪДЪРЖАНИЕ: %s (%s)",
                    source, "; ".join(verdict["evidence"]),
                )

        return {"found": found, "details": details}

    def _list_supported_files(self) -> list[Path]:
        """List supported files in project dir (full recursive scan).

        Excludes files inside the ``converted/`` subfolder.
        """
        if not self.base_path:
            return []

        converted_dir = self.base_path / "converted"
        files: list[Path] = []

        for f in sorted(self.base_path.rglob("*")):
            if not f.is_file():
                continue
            if converted_dir in f.parents:
                continue
            if f.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            if f.name.startswith(("~$", "~_")):
                continue
            files.append(f)

        return files

    def list_files(self) -> list[dict]:
        """List supported project files with metadata.

        Returns:
            List of dicts with name, path, size_kb, extension.
        """
        result = []
        for f in self._list_supported_files():
            stat = f.stat()
            result.append({
                "name": f.name,
                "path": str(f),
                "size_kb": round(stat.st_size / 1024, 1),
                "extension": f.suffix.lower(),
            })
        return result

    def get_supported_files(self) -> list[dict]:
        """Alias for list_files (backward compat)."""
        return self.list_files()

    def get_project_summary(self) -> dict:
        """Get a summary of the project directory.

        Returns:
            Dict with total_files, total_size_kb, by_type, supported_files.
        """
        files = self.list_files()
        if not files:
            return {
                "total_files": 0,
                "total_size_kb": 0,
                "by_type": {},
                "supported_files": 0,
            }

        by_type: dict[str, int] = {}
        total_size = 0.0
        for f in files:
            ext = f["extension"] or "(other)"
            by_type[ext] = by_type.get(ext, 0) + 1
            total_size += f["size_kb"]

        return {
            "total_files": len(files),
            "total_size_kb": round(total_size, 1),
            "by_type": by_type,
            "supported_files": len(files),
        }

    # ------------------------------------------------------------------
    # Manifest management
    # ------------------------------------------------------------------

    def _manifest_path(self) -> Path:
        assert self.base_path is not None
        return self.base_path / "converted" / "_manifest.json"

    def _load_manifest(self) -> None:
        mp = self._manifest_path()
        if mp.exists():
            try:
                self._manifest = json.loads(mp.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                self._manifest = {}
        else:
            self._manifest = {}

    def _save_manifest(self) -> None:
        assert self.base_path is not None
        converted_dir = self.base_path / "converted"
        converted_dir.mkdir(exist_ok=True)

        self._manifest.setdefault("project_path", str(self.base_path))
        self._manifest.setdefault("created", _now_iso())
        self._manifest["last_updated"] = _now_iso()
        self._manifest["app_version"] = APP_VERSION

        # Recompute stats
        files_info: dict = self._manifest.get("files", {})
        ok = sum(1 for v in files_info.values() if v.get("status") == "ok")
        ocr = sum(
            1
            for v in files_info.values()
            if v.get("conversion_method") == "ocr_vision"
        )
        failed = sum(1 for v in files_info.values() if v.get("status") == "error")
        self._manifest["stats"] = {
            "total_files": len(files_info),
            "converted_ok": ok,
            "converted_ocr": ocr,
            "failed": failed,
        }

        self._manifest_path().write_text(
            json.dumps(self._manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    # ------------------------------------------------------------------
    # Conversion status
    # ------------------------------------------------------------------

    def is_conversion_needed(self) -> bool:
        """Check if there are unconverted or changed files.

        Returns:
            True if any files need conversion.
        """
        status = self.get_conversion_status()
        return (status["pending"] + status["changed"]) > 0

    def get_conversion_status(self) -> dict:
        """Compare original files against the manifest.

        Returns:
            Dict with total, converted, pending, changed, failed, method_summary, files (list of details).
        """
        supported = self._list_supported_files()
        manifest_files: dict = self._manifest.get("files", {})
        details: list[dict] = []
        converted = 0
        pending = 0
        changed = 0

        for fp in supported:
            name = fp.name
            stat = fp.stat()
            entry = manifest_files.get(name)

            if entry is None:
                details.append({"name": name, "status": "pending"})
                pending += 1
            elif entry.get("status") == "error":
                details.append({"name": name, "status": "pending"})
                pending += 1
            elif (
                entry.get("original_size") != stat.st_size
                or entry.get("original_modified")
                != datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc)
                .replace(tzinfo=None)
                .isoformat(timespec="seconds")
            ):
                details.append({"name": name, "status": "changed"})
                changed += 1
            else:
                details.append({
                    "name": name,
                    "status": "converted",
                    "method": entry.get("conversion_method", ""),
                })
                converted += 1

        # Build method summary from manifest
        method_counts: dict[str, int] = {}
        for name, entry in manifest_files.items():
            if entry.get("status") == "ok":
                method = entry.get("conversion_method", "unknown")
                method_counts[method] = method_counts.get(method, 0) + 1

        return {
            "total": len(supported),
            "converted": converted,
            "pending": pending,
            "changed": changed,
            "method_summary": method_counts,
            "files": details,
        }

    # ------------------------------------------------------------------
    # File classification (pre-conversion check)
    # ------------------------------------------------------------------

    # Имена, които издават ТАБЛИЦА С КОЛИЧЕСТВА.
    #
    # 21.08.2026: КСС НЕ Е задължителна за график.  Проверено с прогон — шест
    # реда „мрежа, дължина, диаметър, материал" дават чист, валиден и
    # експортируем график, без нито един КСС файл.  Нужни са КОЛИЧЕСТВА, а
    # количествено-стойностната сметка е само един от възможните им носители:
    # тя носи и фасонни части, и единични цени, и суми, а нищо от това не
    # влиза в срока.
    #
    # Затова тук се разпознава и обикновената таблица с дължини.  Иначе такъв
    # файл минаваше за „непознат", целият пакет се конвертираше напразно и чак
    # тогава количествата се намираха по съдържание.
    _REQUIRED_KEYWORDS: frozenset[str] = frozenset({
        "ксс", "кс ", "количествен", "количества", "сметка", "bill", "boq",
        "дължини", "дължина", "ведомост", "quantities",
    })
    # „спецификация" НЕ е тук нарочно: техническата спецификация
    # е друг документ и с нея падаше в „задължителни" вместо в „полезни".

    # Маркери на количествена таблица В СЪДЪРЖАНИЕТО на документа.
    #
    # BACKLOG т.1: класификацията решаваше само по ИМЕТО на файла.  Файл
    # „Техническо предложение.pdf" с таблицата с количества вътре се
    # отхвърляше като незадължителен и генерирането се блокираше — при
    # налични количества.  В българската практика количествата често са
    # приложение към предложението или целият пакет е един PDF.
    #
    # Търсят се заглавия на колони, които се срещат заедно САМО в
    # количествена сметка.
    _BOQ_COLUMN_MARKERS: tuple[frozenset[str], ...] = (
        frozenset({"мярка", "количество"}),
        frozenset({"ед. мярка", "количество"}),
        frozenset({"мярка", "к-во"}),
        frozenset({"unit", "quantity"}),
        frozenset({"наименование", "мярка", "количество"}),
    )

    # Мерни единици, характерни за строителна количествена сметка.
    _BOQ_UNIT_RE = re.compile(
        r"(?:^|[\s|;,\t])(м\s*[23]?|m\s*[23]?|бр\.?|кг|т|л\.?м\.?)(?=[\s|;,\t]|$)",
        re.IGNORECASE | re.MULTILINE,
    )

    # Колко различни реда с мерна единица правят документа количествена сметка.
    _BOQ_MIN_UNIT_ROWS = 5

    # Keywords that indicate useful-but-optional supporting documents
    _USEFUL_KEYWORDS: frozenset[str] = frozenset({
        "технич", "задание", "спецификац", "договор", "проект",
        "обяснителн", "записка", "пояснителн", "техническо",
    })

    # Keywords that indicate a situation / site-plan drawing (трасировъчен план)
    # These files contain street/quarter names as visual labels — ground-truth for locations.
    _SITUATION_KEYWORDS: frozenset[str] = frozenset({
        "ситуация", "ситуат", "трасе", "трасировъч", "situation", "site plan",
        "генерален план", "ген.план", "генплан",
    })

    def classify_files(self, ai_processor: Any | None = None) -> dict:
        """Classify project files as required, useful, situation, or unknown.

        Step 1: keyword match on filename (free, instant).
        Step 2: if no required file found and ai_processor available,
                ask DeepSeek to classify by filename (cheap fallback).

        Returns:
            Dict with keys:
                required        (list[str])   — таблици с количества (КСС или списък с дължини)
                useful          (list[str])   — tech specs, contracts, etc.
                situation       (list[str])   — site plan / трасировъчен план filenames
                situation_paths (list[str])   — full absolute paths to situation files
                unknown         (list[str])   — unrecognised files
                can_proceed     (bool)        — True при поне една таблица с количества
                ai_used         (bool)        — True if AI fallback was triggered
        """
        files = self._list_supported_files()

        required: list[str] = []
        useful: list[str] = []
        situation: list[str] = []
        situation_paths: list[str] = []
        unknown: list[str] = []

        for fp in files:
            name = fp.name
            lower = name.lower()
            if any(kw in lower for kw in self._SITUATION_KEYWORDS):
                situation.append(name)
                situation_paths.append(str(fp))
            elif any(kw in lower for kw in self._REQUIRED_KEYWORDS):
                required.append(name)
            elif any(kw in lower for kw in self._USEFUL_KEYWORDS):
                useful.append(name)
            else:
                unknown.append(name)

        # If keyword match already found a required file — done.
        if required:
            return {
                "required": required,
                "useful": useful,
                "situation": situation,
                "situation_paths": situation_paths,
                "unknown": unknown,
                "can_proceed": True,
                "ai_used": False,
            }

        # Fallback: ask AI to classify by filename only.
        if ai_processor is not None and hasattr(ai_processor, "router") and ai_processor.router:
            try:
                names = [fp.name for fp in files]
                file_list = "\n".join(f"- {n}" for n in names)
                messages = [{
                    "role": "user",
                    "content": (
                        "Класифицирай следните файлове от строителен проект. "
                        "За всеки файл посочи категорията му:\n"
                        "  required  — таблица с КОЛИЧЕСТВА: дължини по диаметри, "
                        "бройки съоръжения, площи. КСС е само един от "
                        "възможните ѝ носители; върши работа и обикновен "
                        "списък с дължини\n"
                        "  useful    — техническа спецификация, договор, проект, задание\n"
                        "  situation — ситуация / трасировъчен план / генерален план (чертеж)\n"
                        "  unknown   — всичко останало\n\n"
                        f"Файлове:\n{file_list}\n\n"
                        "Отговори само с валиден JSON:\n"
                        '{"required": [...], "useful": [...], "situation": [...], "unknown": [...]}'
                    ),
                }]
                file_class_prompt = (
                    "Ти си асистент за класификация на строителни документи. "
                    "Задачата ти е да разпознаеш ролята на всеки файл в тендерна "
                    "документация за ВиК инфраструктурен проект. "
                    "Отговаряй САМО с валиден JSON — без обяснения, без markdown."
                )
                result = ai_processor.router.chat(messages, file_class_prompt)
                classified = ai_processor.router.parse_json_response(result.get("content", "{}"))
                ai_required = classified.get("required", [])
                ai_useful = classified.get("useful", [])
                ai_situation = classified.get("situation", [])
                ai_unknown = classified.get("unknown", [])
                # Resolve full paths for AI-detected situation files
                name_to_path = {fp.name: str(fp) for fp in files}
                ai_situation_paths = [name_to_path[n] for n in ai_situation if n in name_to_path]
                return {
                    "required": ai_required,
                    "useful": ai_useful,
                    "situation": ai_situation,
                    "situation_paths": ai_situation_paths,
                    "unknown": ai_unknown,
                    "can_proceed": len(ai_required) > 0,
                    "ai_used": True,
                }
            except Exception:
                logger.warning("AI file classification failed, proceeding with unknown classification.")

        # No AI available and no keyword match — cannot determine required files.
        return {
            "required": [],
            "useful": useful,
            "situation": situation,
            "situation_paths": situation_paths,
            "unknown": unknown + required,  # required is empty here, unknown gets everything
            "can_proceed": False,
            "ai_used": False,
        }

    # ------------------------------------------------------------------
    # Batch conversion
    # ------------------------------------------------------------------

    def convert_all(
        self,
        ai_processor: Any | None = None,
        progress_callback: Callable[[int, int, str, str], None] | None = None,
        force: bool = False,
    ) -> dict:
        """Convert all pending/changed files.

        Args:
            ai_processor: Optional AIProcessor for OCR on scanned PDFs.
            progress_callback: Called with (current, total, filename, status_emoji).
            force: If True, re-convert ALL files regardless of cache.

        Returns:
            Dict with converted, skipped, failed counts and errors list.
        """
        supported = self._list_supported_files()
        status = self.get_conversion_status()
        file_statuses = {d["name"]: d["status"] for d in status["files"]}

        converted = 0
        skipped = 0
        failed = 0
        errors: list[str] = []
        results: list[dict] = []

        for i, fp in enumerate(supported):
            fname = fp.name
            needs_work = force or file_statuses.get(fname) != "converted"

            if not needs_work:
                skipped += 1
                if progress_callback:
                    progress_callback(i + 1, len(supported), fname, "skip")
                results.append({"file": fname, "action": "skipped"})
                continue

            if progress_callback:
                progress_callback(i + 1, len(supported), fname, "working")

            try:
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
                    future = ex.submit(self.convert_single_file, str(fp), ai_processor)
                    try:
                        result = future.result(timeout=_FILE_CONVERT_TIMEOUT)
                    except concurrent.futures.TimeoutError:
                        result = {"status": "error", "error": f"Timeout ({_FILE_CONVERT_TIMEOUT}s)"}
                if result["status"] == "ok":
                    converted += 1
                    results.append({
                        "file": fname,
                        "action": "converted",
                        "method": result.get("method", ""),
                        "detail": result.get("detail", ""),
                    })
                    if progress_callback:
                        progress_callback(i + 1, len(supported), fname, "done")
                else:
                    failed += 1
                    errors.append(f"{fname}: {result.get('error', 'unknown')}")
                    results.append({"file": fname, "action": "failed", "error": result.get("error")})
                    if progress_callback:
                        progress_callback(i + 1, len(supported), fname, "error")
            except Exception as exc:
                failed += 1
                errors.append(f"{fname}: {exc}")
                results.append({"file": fname, "action": "failed", "error": str(exc)})
                logger.exception("Conversion failed for %s", fname)
                if progress_callback:
                    progress_callback(i + 1, len(supported), fname, "error")

        return {
            "converted": converted,
            "skipped": skipped,
            "failed": failed,
            "errors": errors,
            "results": results,
        }

    # ------------------------------------------------------------------
    # Single file conversion
    # ------------------------------------------------------------------

    def convert_single_file(
        self, filepath: str, ai_processor: Any | None = None
    ) -> dict:
        """Convert a single file based on its extension.

        Args:
            filepath: Absolute path to the file.
            ai_processor: Optional AIProcessor for OCR.

        Returns:
            Dict with status, output_file, method, detail.
        """
        fp = Path(filepath)
        ext = fp.suffix.lower()

        converters = {
            ".pdf": self._convert_pdf,
            ".xlsx": self._convert_excel,
            ".xls": self._convert_excel,
            ".docx": self._convert_docx,
            ".csv": self._convert_csv,
            ".json": self._copy_json_txt,
            ".txt": self._copy_json_txt,
        }

        converter = converters.get(ext)
        if not converter:
            return {"status": "error", "error": f"Unsupported extension: {ext}"}

        try:
            if ext == ".pdf":
                result = converter(str(fp), ai_processor)
            else:
                result = converter(str(fp))
        except Exception as exc:
            logger.exception("Converter error for %s", fp.name)
            return {"status": "error", "error": str(exc)}

        if result.get("status") != "ok":
            return result

        # Write converted JSON
        assert self.base_path is not None
        converted_dir = self.base_path / "converted"
        converted_dir.mkdir(exist_ok=True)
        out_name = fp.stem + ".json"
        out_path = converted_dir / out_name

        out_path.write_text(
            json.dumps(result["data"], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        # Update manifest entry
        stat = fp.stat()
        mtime = (
            datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc)
            .replace(tzinfo=None)
            .isoformat(timespec="seconds")
        )
        self._manifest.setdefault("files", {})[fp.name] = {
            "original_size": stat.st_size,
            "original_modified": mtime,
            "converted_file": f"converted/{out_name}",
            "converted_size": out_path.stat().st_size,
            "conversion_method": result.get("method", ""),
            "conversion_date": _now_iso(),
            "status": "ok",
            "pages_or_rows": result.get("pages_or_rows", 0),
        }
        self._save_manifest()

        logger.info(
            "Converted %s -> %s (%s)", fp.name, out_name, result.get("method")
        )

        return {
            "status": "ok",
            "output_file": str(out_path),
            "method": result.get("method", ""),
            "detail": result.get("detail", ""),
        }

    # ------------------------------------------------------------------
    # PDF conversion
    # ------------------------------------------------------------------

    def _convert_pdf(
        self, filepath: str, ai_processor: Any | None = None
    ) -> dict:
        """Convert a PDF file to structured JSON.

        Strategy (optimized for speed and cost) — решението е ЗА ВСЯКА
        СТРАНИЦА поотделно, не по средно аритметично на документа:
        1. Extract text with PyMuPDF/fitz (best quality, local, free)
        2. Класифицирай всяка страница (`_classify_pdf_pages`)
        3. Сканираните страници -> OCR САМО тях, останалите вървят с fitz текста
        4. Ако няма сканирани, но текстът е рядък -> reformat през DeepSeek
        5. Иначе -> директно fitz текста

        ЗАЩО не средно аритметично (поправено 2026-07-22): документ с 27
        текстови страници и 3 сканирани чертежа дава високо средно и трите
        чертежа минаваха като празни — тиха загуба на данни.  Обратно,
        сканирана страница с тънък текстов слой (10-49 знака от печат или
        колонтитул) попадаше в „reformat", където vision изобщо не се вика,
        и AI-ят преформатираше боклук в убедително изглеждащ боклук.
        """
        import fitz  # PyMuPDF — much better than PyPDF2

        doc = fitz.open(filepath)
        page_info = self._classify_pdf_pages(doc)
        num_pages = len(doc) or 1
        doc.close()

        pages_text = [{"page": p["page"], "text": p["text"]} for p in page_info]
        total_chars = sum(p["chars"] for p in page_info)
        avg_chars = total_chars / num_pages
        scanned = [p for p in page_info if p["kind"] == "scanned"]
        thin = [p for p in page_info if p["kind"] == "thin"]

        source_name = Path(filepath).name

        # --- СКАНИРАНИ СТРАНИЦИ — OCR само тях, останалите остават с fitz ---
        if scanned and ai_processor is not None and hasattr(ai_processor, "ocr_pdf"):
            merged = self._ocr_scanned_pages(
                filepath, source_name, pages_text, scanned, num_pages, ai_processor
            )
            if merged is not None:
                return merged

        if scanned:
            # Няма API за OCR — не се преструвай, че документът е прочетен.
            logger.warning(
                "%s: %d сканирани страници останаха без OCR (няма AI processor).",
                source_name, len(scanned),
            )

        # --- ВСИЧКИ СТРАНИЦИ С ТЕКСТ — използвай директно ---
        if not scanned and not thin:
            full_text = "\n\n".join(p["text"] for p in pages_text if p["text"])
            data = {
                "source_file": source_name,
                "type": "pdf",
                "extraction_method": "fitz_text",
                "pages": num_pages,
                "content": pages_text,
                "full_text": full_text,
            }
            return {
                "status": "ok",
                "data": data,
                "method": "fitz_text",
                "detail": f"{num_pages} стр., {avg_chars:.0f} симв/стр",
                "pages_or_rows": num_pages,
            }

        # --- РЯДЪК ТЕКСТ без сканирани страници — reformat през DeepSeek ---
        if not scanned:
            raw_text = "\n\n".join(p["text"] for p in pages_text if p["text"])

            if ai_processor is not None and hasattr(ai_processor, "reformat_text"):
                try:
                    reformatted = ai_processor.reformat_text(raw_text, source_name)
                    if reformatted.get("status") == "ok":
                        data = {
                            "source_file": source_name,
                            "type": "pdf",
                            "extraction_method": "fitz_reformat",
                            "pages": num_pages,
                            "content": pages_text,
                            "full_text": reformatted["text"],
                        }
                        return {
                            "status": "ok",
                            "data": data,
                            "method": "fitz_reformat",
                            "detail": f"{num_pages} стр., преформатиран (DeepSeek)",
                            "pages_or_rows": num_pages,
                        }
                except Exception as exc:
                    logger.warning("Reformat failed for %s: %s", source_name, exc)

            # Reformat failed or no API — save raw partial text
            data = {
                "source_file": source_name,
                "type": "pdf",
                "extraction_method": "fitz_partial",
                "pages": num_pages,
                "content": pages_text,
                "full_text": raw_text,
            }
            return {
                "status": "ok",
                "data": data,
                "method": "fitz_partial",
                "detail": f"{num_pages} стр., частичен текст ({avg_chars:.0f} симв/стр)",
                "pages_or_rows": num_pages,
            }

        # --- Сканирани страници, но OCR не е бил възможен ---
        full_text = "\n\n".join(p["text"] for p in pages_text if p["text"])
        data = {
            "source_file": source_name,
            "type": "pdf",
            "extraction_method": "no_text",
            "pages": num_pages,
            "scanned_pages": [p["page"] for p in scanned],
            "content": pages_text,
            "full_text": full_text,
        }
        return {
            "status": "ok",
            "data": data,
            "method": "no_text",
            "detail": (
                f"{num_pages} стр., {len(scanned)} сканирани БЕЗ OCR "
                f"(нужен API) — стр. {', '.join(str(p['page']) for p in scanned[:5])}"
            ),
            "pages_or_rows": num_pages,
        }

    # ------------------------------------------------------------------
    # PDF page classification (P1)
    # ------------------------------------------------------------------

    @staticmethod
    def _classify_pdf_pages(doc: Any) -> list[dict]:
        """Класифицирай всяка страница като text / thin / scanned / empty.

        Решаващият сигнал за „сканирана" НЕ е броят знаци, а голямо
        изображение, покриващо страницата.  Титулна страница с 30 знака и
        без картинка е просто рядка (thin) — там OCR няма какво да добави.
        Сканиран чертеж с 30 знака от печат е СКАНИРАН и иска vision.

        Args:
            doc: Отворен fitz документ.

        Returns:
            Списък от {page, text, chars, image_coverage, kind}.
        """
        pages: list[dict] = []

        for i, page in enumerate(doc):
            text = page.get_text().strip()
            chars = len(text)
            coverage = FileManager._image_coverage(page)

            if coverage >= _SCANNED_IMAGE_COVERAGE and chars < _MIN_CHARS_PER_PAGE:
                kind = "scanned"
            elif chars >= _MIN_CHARS_PER_PAGE:
                kind = "text"
            elif chars >= _THIN_TEXT_CHARS:
                kind = "thin"
            else:
                # Без текст и без голямо изображение — празна страница.
                # OCR няма какво да извлече от нея.
                kind = "empty"

            pages.append({
                "page": i + 1,
                "text": text,
                "chars": chars,
                "image_coverage": round(coverage, 3),
                "kind": kind,
            })

        return pages

    @staticmethod
    def _image_coverage(page: Any) -> float:
        """Каква част от страницата е покрита с изображения (0.0–1.0).

        Използва се за разпознаване на сканирани страници.  При грешка или
        липсваща поддръжка връща 0.0 — тогава класификацията пада обратно
        към броя знаци, т.е. към старото поведение.
        """
        try:
            page_rect = page.rect
            page_area = float(page_rect.width) * float(page_rect.height)
            if page_area <= 0:
                return 0.0

            covered = 0.0
            for info in page.get_image_info():
                bbox = info.get("bbox")
                if not bbox or len(bbox) != 4:
                    continue
                width = abs(float(bbox[2]) - float(bbox[0]))
                height = abs(float(bbox[3]) - float(bbox[1]))
                covered += width * height

            return min(covered / page_area, 1.0)
        except Exception as exc:
            logger.debug("Не мога да измеря покритието с изображения: %s", exc)
            return 0.0

    @staticmethod
    def _ocr_scanned_pages(
        filepath: str,
        source_name: str,
        pages_text: list[dict],
        scanned: list[dict],
        num_pages: int,
        ai_processor: Any,
    ) -> dict | None:
        """OCR само сканираните страници и слей резултата с fitz текста.

        Returns:
            Готов conversion резултат, или None ако OCR-ът се провали
            (тогава извикващият продължава по другите пътища).
        """
        indices = [p["page"] - 1 for p in scanned]
        try:
            ocr_result = ai_processor.ocr_pdf(filepath, pages=indices)
        except TypeError:
            # По-стар ai_processor без параметър pages — OCR на целия документ.
            try:
                ocr_result = ai_processor.ocr_pdf(filepath)
            except Exception:
                logger.exception("OCR failed for %s", source_name)
                return None
        except Exception:
            logger.exception("OCR failed for %s", source_name)
            return None

        if ocr_result.get("status") != "ok":
            logger.warning(
                "OCR за %s върна грешка: %s",
                source_name, ocr_result.get("error", "неизвестна"),
            )
            return None

        # Слей: OCR текстът замества само страниците, които са били сканирани.
        ocr_by_page = {
            p.get("page"): p.get("text", "")
            for p in ocr_result.get("data", {}).get("content", [])
        }
        merged = [
            {"page": p["page"], "text": ocr_by_page.get(p["page"], p["text"])}
            for p in pages_text
        ]
        full_text = "\n\n".join(p["text"] for p in merged if p["text"])

        all_scanned = len(scanned) == num_pages
        method = "ocr_vision" if all_scanned else "fitz_ocr_hybrid"
        data = {
            "source_file": source_name,
            "type": "pdf",
            "extraction_method": method,
            "pages": num_pages,
            "ocr_pages": [p["page"] for p in scanned],
            "content": merged,
            "full_text": full_text,
        }
        detail = (
            f"OCR {num_pages} стр."
            if all_scanned
            else f"{num_pages} стр., OCR на {len(scanned)} сканирани"
        )
        return {
            "status": "ok",
            "data": data,
            "method": method,
            "detail": detail,
            "pages_or_rows": num_pages,
        }

    # ------------------------------------------------------------------
    # Excel conversion
    # ------------------------------------------------------------------

    def _convert_excel(self, filepath: str) -> dict:
        """Convert an Excel file (.xlsx/.xls) to structured JSON.

        Handles merged cells by propagating the merged value.
        """
        fp_obj = Path(filepath)
        if fp_obj.suffix.lower() == ".xls":
            import xlrd
            return self._convert_excel_xls(filepath)

        from openpyxl import load_workbook

        wb = load_workbook(filepath, data_only=True)
        sheets: list[dict] = []
        total_rows = 0

        for ws in wb.worksheets:
            # Build a map of merged cells -> top-left value
            merged_vals: dict[tuple[int, int], Any] = {}
            for merge_range in ws.merged_cells.ranges:
                top_left_val = ws.cell(
                    merge_range.min_row, merge_range.min_col
                ).value
                for row in range(merge_range.min_row, merge_range.max_row + 1):
                    for col in range(merge_range.min_col, merge_range.max_col + 1):
                        merged_vals[(row, col)] = top_left_val

            def _cell_value(row: int, col: int) -> Any:
                if (row, col) in merged_vals:
                    return merged_vals[(row, col)]
                return ws.cell(row, col).value

            # Detect header row — SEMANTIC scoring (одит v11 #4).
            #
            # Български КСС често имат слят заглавен ред („ВиК инфраструктура…")
            # ИЛИ ред „Обект | Реконструкция" най-отгоре.  Правилото „първи ред с
            # ≥2 различни стойности" ги избираше погрешно за хедър, а истинският
            # ред (№ | Наименование | Мярка | Количество) ставаше данни.
            #
            # Затова: всеки кандидат-ред се ОЦЕНЯВА по колко ИЗВЕСТНИ имена на
            # колони съдържа; печели редът с най-много.  Само ако никой ред няма
            # известни имена — fallback към „≥2 различни стойности".
            header_keywords = (
                "наименование", "описание", "дейност", "позиция", "мрежа",
                "мярка", "количество", "дължина", "к-во", "кол-во", "№", "код",
                "диаметър", "цена", "unit", "quantity", "description", "item",
            )

            def _header_score(r: int) -> int:
                score = 0
                for c in range(1, (ws.max_column or 1) + 1):
                    v = _cell_value(r, c)
                    low = str(v or "").strip().lower()
                    if low and any(k in low for k in header_keywords):
                        score += 1
                return score

            scan_to = min(ws.max_row or 1, 25) + 1
            scored = [(_header_score(r), -r, r) for r in range(1, scan_to)]
            best_score, _, best_r = max(scored) if scored else (0, 0, 1)
            if best_score >= 2:
                header_row_idx = best_r
            else:
                # Fallback: първи ред с ≥2 РАЗЛИЧНИ стойности (слят заглавен
                # ред има 1 различна стойност и се прескача).
                header_row_idx = 1
                for r in range(1, scan_to):
                    distinct = {
                        str(_cell_value(r, c)).strip()
                        for c in range(1, (ws.max_column or 1) + 1)
                        if _cell_value(r, c) is not None and str(_cell_value(r, c)).strip()
                    }
                    if len(distinct) >= 2:
                        header_row_idx = r
                        break

            # УНИКАЛНИ имена на колони (одит v12 #4): при двуредови/слети
            # headers две колони често имат едно и също име („Количество" за
            # „Дължина" и за „Брой").  В dict еднаквите ключове се презаписват →
            # губи се стойност.  Затова дубликатите получават суфикс.
            headers = []
            _seen: dict[str, int] = {}
            for c in range(1, (ws.max_column or 1) + 1):
                h = str(_cell_value(header_row_idx, c) or f"Col{c}").strip()
                if h in _seen:
                    _seen[h] += 1
                    h = f"{h} ({_seen[h]})"
                else:
                    _seen[h] = 1
                headers.append(h)

            rows: list[dict] = []
            for r in range(header_row_idx + 1, (ws.max_row or 0) + 1):
                row_data: dict[str, Any] = {}
                all_empty = True
                for ci, h in enumerate(headers, start=1):
                    val = _cell_value(r, ci)
                    if val is not None:
                        all_empty = False
                    row_data[h] = _serialize_value(val)
                if not all_empty:
                    # Реалният Excel ред (одит v11 #2): празните редове се
                    # пропускат, затова индексът в списъка НЕ съвпада с реда във
                    # файла.  provenance-цитатът трябва да сочи истинския ред.
                    row_data["__excel_row__"] = r
                    rows.append(row_data)

            total_rows += len(rows)
            sheets.append({
                "name": ws.title,
                "headers": headers,
                "rows": rows,
                "row_count": len(rows),
            })

        wb.close()

        data = {
            "source_file": Path(filepath).name,
            "type": "excel",
            "sheets": sheets,
        }

        sheet_summary = ", ".join(
            f"{s['name']}({s['row_count']})" for s in sheets
        )
        return {
            "status": "ok",
            "data": data,
            "method": "openpyxl",
            "detail": f"{len(sheets)} листа, {total_rows} реда",
            "pages_or_rows": total_rows,
        }

    def _convert_excel_xls(self, filepath: str) -> dict:
        """Convert legacy .xls file using xlrd."""
        import xlrd

        wb = xlrd.open_workbook(filepath)
        sheets: list[dict] = []
        total_rows = 0

        for ws in wb.sheets():
            if ws.nrows == 0:
                continue
            headers = [str(ws.cell_value(0, c) or f"Col{c+1}").strip() for c in range(ws.ncols)]
            rows: list[dict] = []
            for r in range(1, ws.nrows):
                row_data = {headers[c]: ws.cell_value(r, c) for c in range(ws.ncols)}
                if any(v for v in row_data.values() if v != ""):
                    rows.append(row_data)
            total_rows += len(rows)
            sheets.append({"name": ws.name, "headers": headers, "rows": rows, "row_count": len(rows)})

        data = {"source_file": Path(filepath).name, "type": "excel", "sheets": sheets}
        return {
            "status": "ok",
            "data": data,
            "method": "xlrd",
            "detail": f"{len(sheets)} листа, {total_rows} реда",
            "pages_or_rows": total_rows,
        }

    # ------------------------------------------------------------------
    # DOCX conversion
    # ------------------------------------------------------------------

    def _convert_docx(self, filepath: str) -> dict:
        """Convert a DOCX file to structured JSON."""
        from docx import Document

        doc = Document(filepath)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables: list[dict] = []
        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])

            if rows_data:
                headers = rows_data[0]
                data_rows = [
                    dict(zip(headers, r)) for r in rows_data[1:]
                ]
                tables.append({"headers": headers, "rows": data_rows})

        full_text = "\n".join(paragraphs)

        data = {
            "source_file": Path(filepath).name,
            "type": "docx",
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": full_text,
        }
        return {
            "status": "ok",
            "data": data,
            "method": "python-docx",
            "detail": f"{len(paragraphs)} параграфа, {len(tables)} таблици",
            "pages_or_rows": len(paragraphs),
        }

    # ------------------------------------------------------------------
    # CSV conversion
    # ------------------------------------------------------------------

    def _convert_csv(self, filepath: str) -> dict:
        """Convert a CSV file to JSON, auto-detecting delimiter and encoding."""
        raw = Path(filepath).read_bytes()

        # Try encodings in order of likelihood
        content: str | None = None
        used_encoding = ""
        for enc in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
            try:
                content = raw.decode(enc)
                used_encoding = enc
                break
            except (UnicodeDecodeError, ValueError):
                continue

        if content is None:
            return {"status": "error", "error": "Cannot detect file encoding."}

        # Detect delimiter
        sniffer = csv.Sniffer()
        try:
            sample = content[:4096]
            dialect = sniffer.sniff(sample, delimiters=",;\t|")
            delimiter = dialect.delimiter
        except csv.Error:
            delimiter = "," if "," in content[:1000] else ";"

        reader = csv.reader(io.StringIO(content), delimiter=delimiter)
        all_rows = list(reader)

        if not all_rows:
            return {"status": "error", "error": "Empty CSV file."}

        headers = [h.strip() for h in all_rows[0]]
        rows = [dict(zip(headers, r)) for r in all_rows[1:] if any(c.strip() for c in r)]

        data = {
            "source_file": Path(filepath).name,
            "type": "csv",
            "encoding": used_encoding,
            "delimiter": repr(delimiter),
            "sheets": [{
                "name": "Sheet1",
                "headers": headers,
                "rows": rows,
                "row_count": len(rows),
            }],
        }
        return {
            "status": "ok",
            "data": data,
            "method": "csv",
            "detail": f"{len(rows)} реда ({used_encoding})",
            "pages_or_rows": len(rows),
        }

    # ------------------------------------------------------------------
    # JSON / TXT passthrough
    # ------------------------------------------------------------------

    def _copy_json_txt(self, filepath: str) -> dict:
        """Copy JSON (validated) or TXT files into the converted format."""
        fp = Path(filepath)
        raw = fp.read_bytes()

        # Try encodings
        content: str | None = None
        for enc in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
            try:
                content = raw.decode(enc)
                break
            except (UnicodeDecodeError, ValueError):
                continue

        if content is None:
            return {"status": "error", "error": "Cannot decode file."}

        if fp.suffix.lower() == ".json":
            try:
                parsed = json.loads(content)
            except json.JSONDecodeError as exc:
                return {"status": "error", "error": f"Invalid JSON: {exc}"}
            data = {
                "source_file": fp.name,
                "type": "json",
                "content": parsed,
            }
            method = "json_copy"
            detail = "JSON валидиран"
        else:
            data = {
                "source_file": fp.name,
                "type": "txt",
                "content": content,
                "full_text": content,
            }
            method = "txt_copy"
            detail = f"{len(content)} символа"

        return {
            "status": "ok",
            "data": data,
            "method": method,
            "detail": detail,
            "pages_or_rows": content.count("\n") + 1,
        }

    # ------------------------------------------------------------------
    # Reading converted files
    # ------------------------------------------------------------------

    def get_converted_files(self) -> list[dict]:
        """List all successfully converted files.

        Returns:
            List of dicts with original, converted, type, method, size.
        """
        manifest_files: dict = self._manifest.get("files", {})
        result = []
        for name, entry in manifest_files.items():
            if entry.get("status") != "ok":
                continue
            result.append({
                "original": name,
                "converted": entry.get("converted_file", ""),
                "type": Path(name).suffix.lower(),
                "method": entry.get("conversion_method", ""),
                "size": entry.get("converted_size", 0),
            })
        return result

    def read_converted(self, filename: str) -> dict:
        """Read a converted JSON file by original filename.

        Args:
            filename: Original filename (e.g. 'КСС.xlsx').

        Returns:
            Parsed JSON dict.

        Raises:
            FileNotFoundError: If the converted file doesn't exist.
        """
        assert self.base_path is not None
        out_name = Path(filename).stem + ".json"
        converted_path = self.base_path / "converted" / out_name

        if not converted_path.exists():
            raise FileNotFoundError(
                f"Converted file not found: {converted_path}"
            )

        return json.loads(converted_path.read_text(encoding="utf-8"))

    def get_all_text(self, priority: list[str] | None = None) -> str:
        """Combine text from ALL converted files into one large string.

        Useful for sending to AI for analysis.

        Args:
            priority: Имена на файлове, които да излязат ПЪРВИ.  BACKLOG т.2:
                съдържанието се реже по дължина, а редът досега беше азбучен —
                тоест КСС-то можеше да отпадне, защото се казва „К..." и е
                трето по азбука.  Задължителните документи трябва да са отпред.

        Returns:
            Combined text from all converted documents.
        """
        assert self.base_path is not None
        converted_dir = self.base_path / "converted"
        if not converted_dir.exists():
            return ""

        wanted = {str(p) for p in (priority or [])}

        def _rank(path: Path) -> tuple[int, str]:
            """Приоритетните — първи; останалите по азбучен ред."""
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
                source = data.get("source_file", path.stem)
            except (json.JSONDecodeError, OSError):
                source = path.stem
            return (0 if source in wanted else 1, path.name)

        files = sorted(converted_dir.glob("*.json"), key=_rank)

        parts: list[str] = []
        for jf in files:
            if jf.name == "_manifest.json":
                continue
            try:
                data = json.loads(jf.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                continue

            source = data.get("source_file", jf.stem)
            parts.append(f"=== {source} ===")

            if "full_text" in data and data["full_text"]:
                parts.append(data["full_text"])
            elif "sheets" in data:
                for sheet in data["sheets"]:
                    parts.append(f"--- {sheet.get('name', 'Sheet')} ---")
                    for row in sheet.get("rows", []):
                        parts.append(
                            " | ".join(str(v) for v in row.values())
                        )
            elif "content" in data:
                if isinstance(data["content"], str):
                    parts.append(data["content"])
                elif isinstance(data["content"], list):
                    for item in data["content"]:
                        if isinstance(item, dict) and "text" in item:
                            parts.append(item["text"])

            parts.append("")

        return "\n".join(parts)


# ------------------------------------------------------------------
# Utilities
# ------------------------------------------------------------------

def _now_iso() -> str:
    """Return current UTC time as ISO string (no timezone)."""
    return datetime.now(tz=timezone.utc).replace(tzinfo=None).isoformat(
        timespec="seconds"
    )


def _serialize_value(val: Any) -> Any:
    """Make a cell value JSON-serializable."""
    if val is None:
        return None
    if isinstance(val, datetime):
        return val.isoformat()
    if isinstance(val, (int, float, bool)):
        return val
    return str(val)
